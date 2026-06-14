from __future__ import annotations

from pathlib import Path


def read_document(file_path: str) -> str:
    """
    Read txt, docx, or pdf course material and return plain text.

    Raises ValueError with a user-friendly message when the file cannot be read.
    """
    path = Path(file_path)
    if not path.exists():
        raise ValueError(f"文件不存在：{path}")

    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _read_txt(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    raise ValueError("暂不支持该文件格式，请上传 .txt、.docx 或 .pdf 文件。")


def _read_txt(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("TXT 文件编码无法识别，请尝试保存为 UTF-8 后重新上传。")


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("缺少 python-docx 依赖，无法读取 Word 文档。") from exc

    try:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))
        return "\n".join(paragraphs + table_text)
    except Exception as exc:
        raise ValueError(f"Word 文档读取失败：{exc}") from exc


def _read_pdf(path: Path) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise ValueError("缺少 PyMuPDF 依赖，无法读取 PDF 文件。") from exc

    try:
        with fitz.open(str(path)) as doc:
            return "\n".join(page.get_text().strip() for page in doc if page.get_text().strip())
    except Exception as exc:
        raise ValueError(f"PDF 文件读取失败：{exc}") from exc
