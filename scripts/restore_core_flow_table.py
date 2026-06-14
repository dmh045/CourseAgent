from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"


def final_docx_path() -> Path:
    needle = "\u6700\u7ec8\u7248"
    return next(path for path in OUT_DIR.glob("*.docx") if needle in path.name and not path.name.startswith("~$"))


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def main() -> None:
    docx = final_docx_path()
    doc = Document(docx)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if "文档解析 -> Planner -> 摘要/关键词/导图/PPT/讲稿" in existing:
        print("table already present")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["核心流程", "输入", "处理链路", "输出"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")

    rows = [
        (
            "文档生成",
            "PDF/DOCX/TXT、任务目标、模型配置。",
            "文档解析 -> Planner -> 摘要/关键词/导图/PPT/讲稿 -> PPTX/视频生成 -> 校验。",
            "摘要、关键词、思维导图、PPT、讲稿、视频、字幕、运行报告。",
        ),
        (
            "视频分析",
            "课堂视频、可选 SRT、敏感词或重点需求。",
            "字幕识别/读取 -> 字幕分段 -> 知识点总结 -> 时间戳 marker -> highlight 提炼。",
            "视频分析报告、重点时间戳、highlight、导图、字幕文件。",
        ),
        (
            "历史管理",
            "运行状态、产物路径、错误与耗时。",
            "生成 run_id -> 写 manifest -> UI 查询 -> 下载与展示。",
            "可回溯历史记录、固定展示版本、错误诊断入口。",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    body = doc._body._element
    table_el = table._tbl
    body.remove(table_el)
    heading5 = next(p for p in doc.paragraphs if p.text.strip().startswith("5. 实验"))
    idx = list(body).index(heading5._p)
    body.insert(idx, table_el)
    doc.save(docx)
    print(docx)


if __name__ == "__main__":
    main()
