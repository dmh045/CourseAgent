from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

from build_courseagent_paper import (
    DIAGRAM_PATH,
    add_center,
    add_paragraph,
    configure_styles,
    set_cell_shading,
    set_cell_text,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
SCREENSHOT_DIR = OUT_DIR / "ui_screenshots"


def final_docx_path() -> Path:
    """Overwrite the existing final document instead of creating a new deliverable."""
    for path in OUT_DIR.glob("*.docx"):
        if "最终版" in path.name:
            return path
    return OUT_DIR / "CourseAgent_课程项目小论文_最终版.docx"


DOCX_PATH = final_docx_path()


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def create_architecture_diagram() -> None:
    img = Image.new("RGB", (1920, 1180), "#F6F9FC")
    draw = ImageDraw.Draw(img)
    title = font(50, True)
    subtitle = font(23)
    eyebrow = font(19, True)
    h = font(27, True)
    b = font(20)
    small = font(17)

    def text(x: int, y: int, value: str, fnt, fill="#0F172A") -> None:
        draw.text((x, y), value, font=fnt, fill=fill)

    def wrap_lines(value: str, fnt, width: int) -> list[str]:
        lines: list[str] = []
        line = ""
        for ch in value:
            test = line + ch
            if draw.textlength(test, font=fnt) <= width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = ch
        if line:
            lines.append(line)
        return lines

    def wrapped(x: int, y: int, value: str, fnt, width: int, fill="#334155", gap: int = 8) -> int:
        for line in wrap_lines(value, fnt, width):
            draw.text((x, y), line, font=fnt, fill=fill)
            y += fnt.size + gap
        return y

    def pill(x: int, y: int, value: str, fill: str, outline: str) -> None:
        w = int(draw.textlength(value, font=small)) + 34
        draw.rounded_rectangle((x, y, x + w, y + 38), radius=19, fill=fill, outline=outline, width=2)
        text(x + 17, y + 8, value, small, "#334155")

    def card(x: int, y: int, w: int, hgt: int, fill: str, outline: str, title_text: str, lines: list[str], accent: str) -> None:
        draw.rounded_rectangle((x, y, x + w, y + hgt), radius=26, fill=fill, outline=outline, width=3)
        draw.rounded_rectangle((x, y, x + 11, y + hgt), radius=5, fill=accent)
        text(x + 30, y + 24, title_text, h, "#0F172A")
        yy = y + 74
        for line in lines:
            draw.ellipse((x + 32, yy + 9, x + 40, yy + 17), fill=accent)
            yy = wrapped(x + 54, yy, line, b, w - 82, "#334155", 5)
            yy += 4

    def arrow(start: tuple[int, int], end: tuple[int, int], color="#2563EB", width: int = 5) -> None:
        draw.line((start, end), fill=color, width=width)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) >= abs(ey - sy):
            if ex >= sx:
                pts = [(ex, ey), (ex - 20, ey - 12), (ex - 20, ey + 12)]
            else:
                pts = [(ex, ey), (ex + 20, ey - 12), (ex + 20, ey + 12)]
        else:
            if ey >= sy:
                pts = [(ex, ey), (ex - 12, ey - 20), (ex + 12, ey - 20)]
            else:
                pts = [(ex, ey), (ex - 12, ey + 20), (ex + 12, ey + 20)]
        draw.polygon(pts, fill=color)

    text(88, 58, "CourseAgent 架构与数据流设计", title, "#0B2545")
    wrapped(
        90,
        122,
        "面向课程资料生成与课堂视频分析的智能工作台：用 Agent 编排大模型推理，用工具链生成确定性产物，用数据治理支撑追踪、复用和上线扩展。",
        subtitle,
        1360,
        "#475569",
        6,
    )
    pill(1510, 65, "论文 / 答辩图", "#EAF3FF", "#B8D3F5")
    pill(1510, 115, "数据流 + 架构层", "#EEF8F1", "#BBDDC7")

    # Main pipeline
    text(90, 210, "01  用户入口与任务输入", eyebrow, "#2563EB")
    text(545, 210, "02  Agent 决策与生成链路", eyebrow, "#059669")
    text(1250, 210, "03  文件产物与展示交付", eyebrow, "#7C3AED")

    card(90, 250, 360, 230, "#FFFFFF", "#BFD7F5", "交互工作台", ["上传课程文档 / 视频", "填写目标与模型配置", "预览、下载、历史查看"], "#3B82F6")
    card(535, 250, 360, 230, "#FFFFFF", "#BDE8CB", "Agent 编排层", ["AgentState 状态机", "Planner 拆解任务", "Executor 调用工具"], "#22C55E")
    card(955, 250, 360, 230, "#FFFFFF", "#F7D8A8", "AI 推理层", ["Prompt 模板库", "LLM 适配与降级", "JSON Schema 结构输出"], "#F59E0B")
    card(1375, 250, 360, 230, "#FFFFFF", "#D8CCFB", "工具执行层", ["文档解析 / PPTX", "Whisper 字幕识别", "MoviePy 视频合成"], "#8B5CF6")

    arrow((450, 365), (535, 365))
    arrow((895, 365), (955, 365))
    arrow((1315, 365), (1375, 365))

    # Quality loop
    draw.rounded_rectangle((535, 520, 1200, 690), radius=24, fill="#F8FAFC", outline="#CAD5E2", width=3)
    text(570, 550, "质量闭环", h, "#0F172A")
    wrapped(
        570,
        592,
        "Verifier 检查产物完整性，Auto-Repair 修复缺失文件；运行详情记录模型、工具、耗时与错误，帮助定位跑题、空泛和 API 失败。",
        b,
        560,
        "#334155",
        9,
    )
    arrow((715, 480), (715, 520), "#16A34A")
    arrow((1060, 520), (1060, 480), "#16A34A")

    card(1375, 540, 360, 210, "#FFFFFF", "#F9C6D0", "展示与交付", ["PPT / 讲稿 / 视频", "字幕与时间戳", "思维导图与运行报告"], "#F43F5E")
    arrow((1555, 480), (1555, 520), "#7C3AED")

    # Data plane
    draw.rounded_rectangle((90, 780, 1735, 1035), radius=30, fill="#EEF6FF", outline="#B8D3F5", width=3)
    text(130, 815, "数据治理与上线扩展层", h, "#0B2545")
    wrapped(
        130,
        857,
        "本地演示阶段用运行批次和 manifest 管理文件；上线后将同一数据模型迁移到对象存储、关系数据库、缓存、向量索引和异步队列。",
        b,
        1210,
        "#334155",
        7,
    )

    data_cards = [
        (130, 935, 275, "原始输入", "文档、视频、字幕、任务目标"),
        (440, 935, 275, "结构化结果", "摘要、关键词、大纲、导图"),
        (750, 935, 275, "最终产物", "PPTX、MP4、SRT、报告"),
        (1060, 935, 275, "元数据索引", "run_id、状态、耗时、错误"),
        (1370, 935, 300, "线上能力", "对象存储、数据库、缓存、队列"),
    ]
    for x, y, w, name, desc in data_cards:
        draw.rounded_rectangle((x, y, x + w, y + 70), radius=18, fill="#FFFFFF", outline="#C8D7E8", width=2)
        text(x + 22, y + 13, name, small, "#0F172A")
        wrapped(x + 22, y + 39, desc, font(16), w - 42, "#475569", 2)
    for start_x in [405, 715, 1025, 1335]:
        arrow((start_x, 970), (start_x + 40, 970), "#64748B", 4)

    # Vertical data arrows from pipeline to data plane
    arrow((270, 480), (270, 780), "#60A5FA", 4)
    arrow((715, 690), (715, 780), "#34D399", 4)
    arrow((1555, 750), (1555, 780), "#A78BFA", 4)

    draw.rounded_rectangle((90, 1070, 1735, 1140), radius=22, fill="#FFFFFF", outline="#CBD5E1", width=2)
    text(125, 1090, "核心设计思想", h, "#0B2545")
    wrapped(
        335,
        1091,
        "把大模型放进可约束、可校验、可回放的工作流：模型负责理解与规划，工具负责生成文件，数据层负责批次化、索引化、缓存和生命周期管理。",
        b,
        1300,
        "#334155",
        4,
    )
    DIAGRAM_PATH.parent.mkdir(exist_ok=True)
    img.save(DIAGRAM_PATH)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths)
    for i, header in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], header, True)
        set_cell_shading(t.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def shot(doc: Document, filename: str, text: str) -> None:
    path = SCREENSHOT_DIR / filename
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.25))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, text)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    for line in dedent(text).strip().splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(31, 77, 120)


def numbered_block(doc: Document, items: list[str]) -> None:
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f"{i}. {item}")
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def build_document() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    create_architecture_diagram()

    doc = Document()
    configure_styles(doc)

    add_center(doc, "基于大语言模型的课程资料智能生成与课堂视频分析系统设计与实现", 20, True, "0B2545")
    add_center(doc, "题目：基于大语言模型的 CourseAgent 课程辅助系统", 12)
    add_center(doc, "学号：______________    姓名：______________", 12)
    add_center(doc, "课程项目小论文", 11, False, "475569")
    doc.add_paragraph()

    doc.add_heading("摘要", level=1)
    add_paragraph(
        doc,
        "本文设计并实现了一个面向课程资料整理、课堂展示准备和课堂视频复习的 CourseAgent 系统。系统并不是简单调用大语言模型生成文本，而是围绕“输入资料—任务规划—工具执行—产物校验—历史归档”的完整链路构建智能工作流。系统采用 Streamlit 构建交互式工作台，使用大语言模型完成内容理解、任务分解和结构化生成，结合文档解析、PPTX 生成、MoviePy 视频合成、本地 Whisper 语音识别、运行记录管理等确定性工具，支持从 PDF、DOCX、TXT、课堂视频和 SRT 字幕中生成摘要、关键词、思维导图、答辩 PPT、逐页讲稿、讲解视频、课堂视频知识点、重点时间戳和 highlight 片段。开发过程中，系统初始产物曾出现摘要空泛、PPT 跑题、关键词碎片化、思维导图层级过浅、讲稿机械朗读等问题。针对这些问题，后续从提示词分层、原文边界约束、JSON Schema 结构化输出、质量校验、失败修复和历史记录等方面进行了迭代，使系统逐步从“功能集合”转向“可执行、可检查、可追溯的课程辅助 Agent”。论文重点分析系统构建思路、技术架构、数据流设计、核心功能实现、实验验证与后续上线扩展方向。",
    )
    add_paragraph(doc, "关键词：CourseAgent；大语言模型；Agent 工作流；提示词工程；数据流架构；课堂视频分析；PPT 自动生成")

    doc.add_heading("1. 开发背景、意义与系统功能", level=1)
    add_paragraph(
        doc,
        "在软件工程、交互设计、人机交互等课程学习中，学生面对的资料往往不是一份规整的讲义，而是一组分散的课件、论文、教材节选、课堂录屏和字幕文件。真正困难的地方并不只是“读完资料”，而是把这些资料重新组织成可以展示、可以讲解、也可以复习的学习产物。人工处理时，用户需要反复阅读材料、提炼知识点、搭建展示逻辑、制作 PPT、撰写讲稿、定位视频重点片段。这个过程耗时长，也很容易因为资料量大、概念抽象或表达目标不清，导致内容遗漏、结构混乱和展示材料质量不稳定。",
    )
    add_paragraph(
        doc,
        "CourseAgent 的开发背景正来源于这一真实学习场景。系统希望解决的不是单一的“让模型回答问题”，而是把课程资料加工过程工程化：用户输入的是资料和目标，系统输出的是一组可交付文件。这样的系统必须同时处理自然语言理解、任务分解、文件生成、视频处理、错误恢复和交互体验等问题。因此，CourseAgent 的建设重点不是把多个功能按钮堆在页面上，而是建立一条能够被用户一键触发、被系统分步执行、被结果校验机制检查的课程资料生产流水线。",
    )
    add_paragraph(doc, "从应用意义看，CourseAgent 的价值主要体现在以下三点：")
    numbered_block(
        doc,
        [
            "效率价值：系统把摘要、提纲、PPT 初稿、讲稿和字幕整理等重复劳动交给自动化流程完成，让学生把更多精力放在理解课程内容和打磨表达上。",
            "结构价值：系统以统一的知识结构驱动多个产物生成，使摘要、导图、PPT、讲稿和视频分析围绕同一主题展开，减少不同材料之间互相割裂的问题。",
            "复盘价值：课堂视频不再只是一个线性播放的长文件，而是被转化为带时间戳、知识点和 highlight 标记的复习索引，用户可以更快回到真正重要的片段。",
        ],
    )
    add_paragraph(doc, "系统功能可以概括为两个主工作流和一个支撑中心：")
    numbered_block(
        doc,
        [
            "课程文档生成工作流：用户上传课程资料并输入目标，系统自动完成文档解析、摘要生成、关键词提取、思维导图生成、PPT 大纲生成、逐页讲稿生成、PPTX 写入、讲解视频合成和产物校验。",
            "课堂视频分析工作流：用户上传课堂视频或字幕，系统完成字幕识别或读取、知识点总结、时间戳标记、highlight 提炼和导图生成。",
            "产物与运行支撑中心：系统保存每次任务的运行记录、产物清单、日志、错误信息和下载入口，避免课堂展示前准备好的版本被后续测试覆盖。",
        ],
    )
    table(
        doc,
        ["功能域", "主要能力", "设计目标"],
        [
            ("课程文档生成", "读取课程 PDF、DOCX、TXT，生成摘要、关键词、思维导图、PPT、讲稿和讲解视频。", "把原始学习资料转化为可展示、可讲解、可下载的成果。"),
            ("课堂视频分析", "识别或读取字幕，提炼知识点、重点时间戳、highlight 片段和复习导图。", "让长视频具有可检索、可复盘、可跳转的知识结构。"),
            ("Agent 工作流", "根据用户目标动态规划工具链，按节点执行并记录中间状态。", "避免固定流程无法适应不同任务，体现系统的智能调度能力。"),
            ("质量闭环", "对模型输出和文件产物进行校验，必要时执行修复或降级。", "降低生成失败、空泛输出、跑题内容和缺失文件的风险。"),
            ("历史记录", "按运行批次保存 manifest、产物路径、耗时、错误和状态。", "支持课堂展示版本固定、结果回溯和后续对比分析。"),
        ],
        [1600, 4500, 3260],
    )

    doc.add_heading("2. 系统主要设计框架", level=1)
    add_paragraph(
        doc,
        "CourseAgent 的系统框架采用“交互工作台 + Agent 编排 + 工具执行 + 数据治理”的架构思想。交互工作台负责让用户以最少步骤完成上传、目标描述、运行和下载；Agent 编排层负责把自然语言目标转化为可执行计划；工具执行层负责将计划中的每一步落成文件、图像、视频或结构化数据；数据治理层负责记录每次运行的输入、输出、状态、耗时和错误。这样的划分避免把大语言模型当作唯一后端，也避免把业务逻辑写成不可追踪的页面脚本。",
    )
    add_paragraph(doc, "在构建思路上，系统并不是把所有逻辑写成一个长函数，而是围绕三个原则组织：")
    numbered_block(
        doc,
        [
            "把课程资料处理抽象为任务状态机。一次运行不是单个函数调用，而是一组状态迁移：初始化输入、读取文档、规划工具、执行工具、校验产物、修复异常、写入报告。这样每一步都有明确输入和输出，UI 可以展示进度，日志也能定位问题。",
            "把不确定的 AI 推理和确定性的文件工具分离。大语言模型适合理解资料、提炼结构和组织表达，但 PPTX 写入、视频合成、字幕保存、文件下载和历史记录必须由程序工具完成，才能保证结果可落地。",
            "把运行结果变成可管理资产。系统通过 run_id、manifest 和产物索引记录每次任务，使生成结果不再是临时文件，而是可以回看、下载、比较和清理的学习资产。",
        ],
    )
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, "图 1  CourseAgent 架构与数据流设计")
    add_paragraph(
        doc,
        "图 1 展示了系统的总体架构。交互层通过 Streamlit 的响应式页面完成文件上传、任务目标输入、API 配置和产物下载。Agent 编排层维护一次任务运行的状态对象，并按照 Planner、Executor、Verifier、Auto-Repair、Report 的顺序推进。AI 推理层封装不同模型供应商的接口，并通过提示词模板、结构化 JSON 输出和质量规则约束模型行为。工具执行层负责文档解析、PPTX 生成、MoviePy 视频合成、Whisper 字幕识别和思维导图渲染。数据治理层则保存运行批次、产物索引、错误日志和缓存。",
    )
    add_paragraph(
        doc,
        "从技术选型看，Streamlit 适合快速构建数据与 AI 工作台，它能够以较低成本实现上传控件、表单、状态提示、下载按钮和多标签页面。Python 生态则适合连接文档、视频和 AI 工具链，例如 python-docx、python-pptx、PyMuPDF、MoviePy、Whisper 等库可以覆盖课程资料处理的主要需求。大语言模型接口采用兼容式适配层，原因是不同 API 在模型名称、Base URL、JSON 输出能力、额度限制和错误格式上存在差异，系统必须把这些差异封装在统一客户端中，避免 UI 和工具模块直接依赖某一个供应商。",
    )
    table(
        doc,
        ["架构层", "关键技术", "建设思路"],
        [
            ("交互层", "Streamlit、表单状态、上传控件、下载按钮、标签页。", "优先保证课堂展示场景下的一键操作和结果可见性，而不是只提供命令行脚本。"),
            ("编排层", "任务状态机、Planner、Executor、Verifier、Auto-Repair。", "把 AI 工作拆成可观察节点，使系统能规划、执行、校验和回放。"),
            ("AI 推理层", "多供应商 LLM 适配、Prompt 模板、JSON Schema、fallback。", "让模型负责高层理解与结构生成，同时用格式约束和质量规则降低不确定性。"),
            ("工具层", "文档解析、PPTX、MoviePy、Whisper、导图渲染。", "把文件生成、视频处理和字幕保存交给确定性工具，减少模型直接生成文件的不可控风险。"),
            ("数据层", "运行批次、manifest、对象文件、缓存、后续可扩展数据库。", "把输入、产物和日志纳入统一数据流，支持追溯、复用、清理和上线扩容。"),
        ],
        [1400, 3300, 4660],
    )
    add_paragraph(
        doc,
        "数据层设计不能只理解为“文件存在哪个本机路径”。在系统设计层面，数据流应分为五类：原始输入、解析中间数据、AI 结构化结果、最终产物和运行元数据。原始输入包括用户上传的文档、视频和字幕；解析中间数据包括文档文本、分段内容、字幕片段和视频时长信息；AI 结构化结果包括摘要、关键词、PPT 大纲、讲稿、思维导图节点和视频 marker；最终产物包括 PPTX、MP4、SRT、Markdown 报告和图片；运行元数据记录任务目标、状态、耗时、错误、模型配置和产物索引。这样划分后，系统即使从本地演示迁移到线上，也可以保持清晰的数据边界。",
    )
    add_paragraph(
        doc,
        "如果系统真正上线，数据层需要从本地文件目录扩展为“对象存储 + 元数据库 + 缓存 + 向量索引 + 异步任务队列”的组合。对象存储用于保存大文件，如视频、PPT、字幕和导图图片；关系型数据库用于保存用户、任务、运行状态、产物清单、权限和配额；Redis 或类似缓存用于保存短期任务状态、进度条和热点中间结果；向量数据库用于保存课程文档分段向量，支持长文档检索增强生成；异步队列用于把视频识别、PPT 生成和视频合成等长耗时任务交给后台 Worker。面对庞大数据时，还需要设置文件生命周期策略，例如临时上传文件定期清理、历史产物按用户选择归档、缓存按最近使用时间淘汰、课堂展示版本单独固定，避免存储无限增长。",
    )
    code_block(
        doc,
        """
        用户上传资料
          -> 输入校验与任务创建
          -> 原始文件写入对象存储 / 临时存储
          -> 解析文本、字幕、视频元数据
          -> LLM 生成结构化 JSON
          -> 工具生成 PPTX、MP4、SRT、导图和报告
          -> Verifier 校验产物完整性
          -> manifest / 数据库记录任务、状态、产物索引
          -> UI 通过 run_id 查询历史记录并提供下载
        """,
    )
    add_paragraph(
        doc,
        "因此，本项目的架构重点并不是“某一层由哪个 .py 文件实现”，而是通过分层和数据流设计解决 AI 应用常见的工程问题：模型输出不稳定、长任务耗时、文件产物容易覆盖、错误不易定位、不同 API 不兼容、视频处理可能拖垮主进程。当前本地版本已经通过运行批次、manifest、缓存、子进程隔离和降级策略解决了部分问题；如果上线，则可以沿着相同的数据模型平滑迁移到数据库、对象存储和异步 Worker 架构。",
    )

    doc.add_heading("3. AI 提示词设计、响应反馈与系统设计思路", level=1)
    add_paragraph(
        doc,
        "提示词设计是 CourseAgent 从“功能集合”走向“智能系统”的关键。项目初期的做法更接近直接让模型生成摘要、PPT 或导图，结果并不理想：上传交互设计课程资料时，PPT 大纲中出现了与课程无关的 Agent 工作流程；关键词只提取到 Design、Mock、the 等碎片词；思维导图只有少量一级节点，不能体现课程知识结构；PPT 内容像模板说明；讲稿只是照着幻灯片逐条朗读。这些问题说明，AI 产物质量差并不一定是模型完全不可用，而是任务边界、输入依据、输出结构和质量标准没有被工程化表达。",
    )
    add_paragraph(
        doc,
        "后续的提示词设计采用“分任务、分角色、分结构、可校验”的策略。Planner 提示词关注任务拆解，要求根据用户目标选择必要工具，而不是固定执行全部功能；摘要提示词关注资料主题、课程概念、结构关系和可用于答辩的表达；关键词提示词强调课程术语、方法名、模型名和流程名，过滤孤立停用词；思维导图提示词要求至少包含主题、核心概念、方法步骤、应用场景和结论层级；PPT 大纲提示词要求每页有标题、核心论点、支撑解释和讲解意图；讲稿提示词要求用教学语言解释知识，而不是重复 PPT 文案；视频分析提示词要求输出时间戳、引用片段、重要性原因和标签。",
    )
    table(
        doc,
        ["初始问题", "根本原因", "改进后的提示词/反馈设计"],
        [
            ("摘要空泛", "只要求概括，没有规定必须绑定原文主题、课程概念和学习目标。", "加入“严格基于资料”“提炼课程结构”“突出可用于展示和复习的价值”等约束。"),
            ("PPT 跑题", "模型把系统背景、Agent 工作流等上下文误当成课程内容。", "加入原文边界和负面约束，明确禁止引入资料之外的系统实现内容。"),
            ("关键词碎片化", "模型直接抽取高频词，缺少领域过滤。", "要求输出概念、方法、模型、流程和术语，并过滤英文停用词与孤立单词。"),
            ("导图简陋", "没有规定层级深度和节点语义。", "要求多层级知识树，并使用质量检查识别 sparse_mindmap。"),
            ("讲稿机械", "讲稿被当成 PPT 文案扩写。", "提示词改为“面向同学讲解”，要求举例、过渡、解释因果关系。"),
            ("视频重点不清", "字幕分析没有规定时间戳、片段和重要性理由。", "要求输出 marker、summary、quote、reason、tags 和 highlight 建议。"),
        ],
        [1650, 3500, 4210],
    )
    add_paragraph(doc, "响应反馈机制体现为四个层次：")
    numbered_block(
        doc,
        [
            "格式反馈：尽量让模型输出 JSON，程序可以解析字段并传给下游工具，避免模型只给一段无法继续加工的散文式回答。",
            "内容反馈：通过规则检查是否出现与资料无关的系统词、节点是否过少、关键词是否过短，从而及时发现跑题和空泛问题。",
            "产物反馈：Verifier 检查计划中需要的 PPT、视频、字幕、导图和报告是否存在且非空，发现缺失后进入修复或失败记录。",
            "用户反馈：UI 通过产物预览、运行详情和错误提示暴露问题，用户可以调整任务目标、模型配置或输入资料后重新运行。",
        ],
    )
    add_paragraph(doc, "从系统思路看，CourseAgent 的智能性主要体现在三个方面：")
    numbered_block(
        doc,
        [
            "任务理解：系统读取用户目标后，判断是否需要摘要、导图、PPT、讲稿、视频或课堂视频分析，而不是所有任务都走同一套固定流程。",
            "工具协同：模型负责生成结构化内容，实际文件由工具完成，既保留模型的理解能力，也降低幻觉和格式错误。",
            "结果治理：系统记录每一步的状态、耗时和错误，方便追踪为什么某次生成结果不理想，也方便后续改进提示词和工具链。",
        ],
    )
    add_paragraph(
        doc,
        "因此，这里的“智能”不是脱离场景的强 AI，而是面向课程资料处理的可控 Agent：它能在明确边界内理解目标、选择工具、产出文件并校验结果。",
    )

    doc.add_heading("4. 核心功能设计：设计图与 UI 功能演示", level=1)
    add_paragraph(
        doc,
        "核心功能围绕两个工作台展开：课程文档生成工作台和课堂视频分析工作台。二者共享 API 配置、运行状态、历史记录和下载能力，但处理逻辑不同。文档工作台的目标是把资料转化为展示材料，因此关注内容结构、PPT 逻辑、讲稿表达和最终文件；视频工作台的目标是把课堂录屏转化为复习索引，因此关注字幕、时间戳、知识点、highlight 和导图。这样的拆分使 UI 与任务模型保持一致，用户不需要在一个混杂页面中寻找功能。",
    )
    shot(doc, "ui_01_document_workspace.png", "图 2 课程文档生成工作台：上传课程资料、填写目标并一键生成摘要、PPT、讲稿和视频")
    add_paragraph(
        doc,
        "图 2 展示了课程文档生成工作台。页面顶部显示当前运行模式、API 状态、模型和输出目录，帮助用户判断系统是否在真实 API、mock 还是降级模式下运行。中间区域将流程显式拆成上传资料、Agent 规划、内容生成、产物校验和下载交付五个步骤，使用户能理解一键生成背后发生了什么。任务目标输入框是系统智能性的入口，用户可以写出“整理成 5 分钟答辩 PPT，并生成摘要、关键词、思维导图、每页讲稿和讲解视频”等自然语言需求，Planner 根据目标选择工具链。",
    )
    shot(doc, "ui_02_video_analysis.png", "图 3 课堂视频分析工作台：上传课堂录屏或 SRT 字幕，生成知识点、导图和重点时间戳")
    add_paragraph(
        doc,
        "图 3 展示了课堂视频分析工作台。该页面不再保留单纯的视频剪辑工具，而是围绕 CourseAgent 的课程场景重新设计：用户上传课堂视频，系统识别或读取字幕，再基于字幕分析知识点和时间戳。如果用户已经有 SRT 字幕，可以直接上传以节省语音识别时间；如果没有字幕，系统调用本地 Whisper 后端。为了避免本地语音识别崩溃导致 Streamlit 主进程断开，系统把识别过程放入子进程，并提供 PyTorch Whisper 备用后端。这一设计体现了“长耗时、易崩溃任务隔离”的工程思想。",
    )
    shot(doc, "ui_03_asset_history.png", "图 4 产物中心：按运行批次保存历史记录，支持展示版本固定与下载")
    shot(doc, "ui_04_run_details.png", "图 5 运行详情：展示 Agent 决策、执行日志、校验结果和错误信息")
    add_paragraph(
        doc,
        "图 4 和图 5 分别展示产物中心和运行详情。产物中心本质上是历史记录页面，每次分析都会形成独立 run_id，并通过 manifest 记录任务类型、源文件、目标、状态、产物路径、耗时和错误信息。这样做可以解决课堂展示中的实际问题：提前生成的文档分析产物不会因为之后分析视频而被覆盖。运行详情则用于解释系统行为，例如 Planner 为什么选择某个工具、Executor 是否调用成功、Verifier 是否发现缺失产物、API 是否返回错误。对 AI 系统来说，运行详情不仅是日志页面，也是调试生成质量的反馈入口。",
    )
    table(
        doc,
        ["核心流程", "输入", "处理链路", "输出"],
        [
            ("文档生成", "PDF/DOCX/TXT、任务目标、模型配置。", "文档解析 -> Planner -> 摘要/关键词/导图/PPT/讲稿 -> PPTX/视频生成 -> 校验。", "摘要、关键词、思维导图、PPT、讲稿、视频、字幕、运行报告。"),
            ("视频分析", "课堂视频、可选 SRT、敏感词或重点需求。", "字幕识别/读取 -> 字幕分段 -> 知识点总结 -> 时间戳 marker -> highlight 提炼。", "视频分析报告、重点时间戳、highlight、导图、字幕文件。"),
            ("历史管理", "运行状态、产物路径、错误与耗时。", "生成 run_id -> 写 manifest -> UI 查询 -> 下载与展示。", "可回溯历史记录、固定展示版本、错误诊断入口。"),
        ],
        [1350, 2300, 3800, 1910],
    )

    doc.add_heading("5. 实验、验证或数据分析与结论", level=1)
    add_paragraph(
        doc,
        "系统验证围绕功能完整性、生成质量、稳定性和交互可用性展开。功能完整性关注计划中的产物是否真的生成；生成质量关注内容是否贴合上传资料、结构是否清晰、PPT 和讲稿是否可用于展示；稳定性关注 API 失败、本地模型崩溃、额度限制和长视频处理时系统是否仍能给出可理解的错误提示；交互可用性关注用户是否能在 UI 中完成上传、运行、预览、下载和历史记录查看。",
    )
    table(
        doc,
        ["验证对象", "验证方式", "结论"],
        [
            ("动态规划", "使用不同目标触发文档生成、思维导图生成和视频分析任务。", "系统能根据目标选择工具链，并为视频或 PPT 任务补齐依赖步骤。"),
            ("文档处理", "上传 KLM 效率模型、交互设计过程等课程 PDF。", "能生成摘要、关键词、导图、PPT、讲稿和报告；质量受模型能力和提示词影响明显。"),
            ("视频分析", "上传课堂视频或 SRT 字幕，检查字幕、marker、highlight 和导图。", "上传 SRT 时最稳定；本地识别可用但耗时较长，需要进度反馈和后台化优化。"),
            ("异常处理", "模拟 OpenAI Key 错误、Gemini 额度耗尽、DeepSeek response_format 不支持、Whisper 子进程异常。", "系统通过错误诊断、接口降级、子进程隔离和备用识别后端提升了可恢复性。"),
            ("自动化测试", "运行项目测试集，覆盖 Planner、Verifier、视频分析和工具依赖。", "最近一次显式测试结果为 18 passed，说明核心模块在测试输入下可正常工作。"),
        ],
        [1600, 4400, 3360],
    )
    add_paragraph(
        doc,
        "从实验现象看，CourseAgent 的效果提升主要来自两个方面：一是提示词和结构化输出的改进，二是工程闭环的完善。初始版本虽然已经能调用模型，但由于任务边界不清，产物容易跑题或空泛；优化后，模型被限制在课程资料范围内，并按固定字段输出，后续工具可以继续加工。另一方面，仅有模型输出并不能保证系统可用，只有加入文件校验、错误日志、历史记录和降级策略后，用户才可以知道哪些内容已经生成、哪些步骤失败、失败原因是什么。",
    )
    add_paragraph(
        doc,
        "结论上，CourseAgent 已经具备课程辅助 Agent 的基本形态：它能够接收用户目标，读取课程资料，规划多步任务，调用不同工具生成学习产物，并记录运行过程。系统的优势在于把课程资料处理从一次性问答变成可重复执行的工作流；不足在于生成质量仍依赖高质量模型，视频处理耗时较长，PPT 视觉设计仍需要更专业的模板系统。对于课堂展示场景，当前系统可以作为辅助生成和初稿整理工具；对于正式上线，还需要继续完善多用户、数据库、异步队列、权限、监控和数据治理能力。",
    )

    doc.add_heading("6. 系统存在问题与改进", level=1)
    add_paragraph(doc, "当前系统的主要问题可以分为四类：")
    numbered_block(
        doc,
        [
            "质量问题：大语言模型对复杂课程资料的理解仍可能不够深入，尤其当资料篇幅较长、概念关系复杂或 API 模型能力较弱时，PPT 和导图仍可能偏概括。",
            "性能问题：视频识别、TTS 语音生成、MoviePy 合成和大文件读写都属于长耗时任务，如果全部在 Streamlit 同步流程中执行，用户会明显感到等待时间过长。",
            "架构问题：当前版本以单机演示为主，虽然已经有运行批次和历史记录，但还没有真正的用户体系、数据库事务、后台队列和分布式 Worker。",
            "体验问题：PPT 的视觉设计、视频时间轴交互、highlight 展示和细粒度进度反馈还可以继续增强。",
        ],
    )
    add_paragraph(doc, "后续改进可以从以下方向展开：")
    numbered_block(
        doc,
        [
            "完善知识增强机制。对于长文档，可以引入分段切片、向量检索和证据引用，使摘要、PPT 和讲稿不只依赖一次性上下文，而是能够回到资料片段中寻找依据。",
            "建立更系统的质量评估机制。系统可以检查 PPT 是否每页有明确论点、讲稿是否重复幻灯片、导图是否达到层级要求、关键词是否为课程概念、视频 marker 是否覆盖关键知识点。",
            "把长耗时任务后台化。视频识别、PPT 生成和视频合成应交给异步队列与 Worker 执行，UI 只负责提交任务、展示进度和接收结果。",
            "继续优化展示体验。后续可以引入更专业的 PPT 模板、视频时间轴标记、highlight 可视化和更直观的历史记录筛选方式，让系统更适合课堂展示。",
        ],
    )
    add_paragraph(
        doc,
        "从上线架构看，需要把本地文件式历史记录升级为数据库驱动的任务系统。用户上传的大文件应进入对象存储，任务元数据进入关系型数据库，后台 Worker 通过队列执行识别、生成和合成任务，UI 只负责提交任务和轮询进度。对于庞大数据，需要设置分层存储和生命周期策略：临时上传文件短期保留，生成产物按用户选择归档，缓存按最近使用时间淘汰，课堂展示版本单独固定，日志和错误信息进入可查询的监控系统。这样才能避免多用户环境下存储无限增长和任务互相影响。",
    )
    add_paragraph(
        doc,
        "从智能性提升看，CourseAgent 可以进一步加入多 Agent 协作思路。例如由资料分析 Agent 负责提炼知识结构，由教学设计 Agent 负责组织讲解逻辑，由视觉设计 Agent 负责 PPT 版式建议，由审查 Agent 负责检查跑题、空泛和证据不足。多个 Agent 不一定都对应独立模型进程，也可以表现为不同的提示词角色和审查阶段。其核心目标是让系统不仅能生成文件，还能像助教一样判断材料是否讲得清楚、结构是否适合课堂展示、重点是否便于复习。",
    )
    add_paragraph(
        doc,
        "综上，CourseAgent 的价值不在于简单堆叠摘要、PPT、视频等功能，而在于把大语言模型、文档解析、视频处理、提示词工程、质量校验和历史记录组织成一个面向课程学习场景的完整工作流。随着数据层、异步任务、检索增强、质量评估和多 Agent 协作继续完善，该系统可以从本地课程项目进一步发展为可上线、可扩展、可持续优化的课程资料智能处理平台。",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CourseAgent 课程项目小论文")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
