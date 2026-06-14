from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"


def final_docx_path() -> Path:
    needle = "\u6700\u7ec8\u7248"
    return next(path for path in OUT_DIR.glob("*.docx") if needle in path.name and not path.name.startswith("~$"))


def style_run(run, size: float = 10.5, bold: bool = False, color: str = "111827") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def blacken_heading_styles(doc: Document) -> None:
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.font.size = Pt(size)


def set_para_body(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.18
    for run in paragraph.runs:
        style_run(run, 10.5, False, "111827")


def insert_heading_before(target, title: str, note: str | None = None) -> None:
    heading = target.insert_paragraph_before(title)
    heading.style = "Heading 2"
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)
    for run in heading.runs:
        style_run(run, 12, True, "000000")
    if note:
        para = target.insert_paragraph_before(note)
        set_para_body(para)


def find_paragraph(doc: Document, prefix: str, start: int = 0):
    for i, paragraph in enumerate(doc.paragraphs[start:], start=start):
        text = paragraph.text.strip()
        if text.startswith(prefix) or prefix in text:
            return i, paragraph
    raise ValueError(f"Cannot find paragraph starting with: {prefix}")


def drawing_before_caption(doc: Document, caption_prefix: str):
    idx, _ = find_paragraph(doc, caption_prefix)
    for j in range(idx - 1, -1, -1):
        if doc.paragraphs[j]._p.xpath(".//w:drawing"):
            return doc.paragraphs[j]
    raise ValueError(f"Cannot find drawing before caption: {caption_prefix}")


def add_once(doc: Document, title: str, anchor_prefix: str, note: str | None = None, before_drawing_caption: bool = False) -> None:
    if any(p.text.strip() == title for p in doc.paragraphs):
        return
    if before_drawing_caption:
        target = drawing_before_caption(doc, anchor_prefix)
    else:
        _, target = find_paragraph(doc, anchor_prefix)
    insert_heading_before(target, title, note)


def main() -> None:
    docx = final_docx_path()
    doc = Document(docx)
    blacken_heading_styles(doc)

    additions = [
        (
            "1.1 场景痛点：从资料堆积到可展示成果",
            "在软件工程、交互设计、人机交互等课程学习中",
            "本节先从真实学习任务出发，说明为什么课程资料处理不是简单阅读，而是一个需要整理、提炼、表达和复盘的连续过程。",
            False,
        ),
        (
            "1.2 系统定位：从聊天问答到学习工作流",
            "CourseAgent 的开发背景正来源于这一真实学习场景",
            "在明确痛点之后，系统定位需要从“让模型回答问题”进一步提升为“让系统完成一组可交付产物”，这也是 CourseAgent 与普通聊天页面的区别。",
            False,
        ),
        (
            "1.3 功能组成：文档生成、视频分析与历史支撑",
            "系统功能可以概括为两个主工作流和一个支撑中心",
            "为了让系统目标更加具体，本节将功能划分为两条主线和一个支撑中心，分别对应资料生成、课堂视频复盘和产物管理。",
            False,
        ),
        (
            "2.1 总体架构：交互、编排、工具与数据治理",
            "CourseAgent 的系统框架采用",
            "前一节说明了系统要解决什么问题，本节进一步说明系统如何组织这些能力，使 UI、Agent、模型、工具和数据层各自承担明确职责。",
            False,
        ),
        (
            "2.2 构建原则：状态机、工具链与可追踪资产",
            "在构建思路上，系统并不是把所有逻辑写成一个长函数",
            "架构不是文件清单，而是一套处理复杂任务的方法。这里重点说明系统为什么采用状态机和工具链，而不是把模型调用和文件生成混在一起。",
            False,
        ),
        (
            "2.3 数据链路：从输入资料到历史下载",
            "在数据层设计层面，数据流分为五类",
            "当系统从本地演示走向真实使用时，最关键的是数据能否被清楚地追踪。本节将数据流拆成几个阶段，说明每一步产生什么、保存什么、供谁继续使用。",
            False,
        ),
        (
            "2.4 上线扩展：面向大文件和多用户的设计",
            "如果系统真正上线，数据层需要从本地文件目录扩展为",
            "本地版本解决的是单人演示和课程项目验证，上线版本还要面对大量文件、并发任务和长期存储，因此需要进一步考虑对象存储、数据库、缓存和异步队列。",
            False,
        ),
        (
            "3.1 初始问题：生成质量为什么不稳定",
            "提示词设计是 CourseAgent 从",
            "前面的架构保证系统能跑起来，但真正决定用户感受的是生成质量。本节先回到开发初期的失败现象，分析为什么最初的摘要、PPT 和导图不够可靠。",
            False,
        ),
        (
            "3.2 提示词重构：分任务、分角色、分结构",
            "后续的提示词设计",
            "针对这些问题，提示词不再写成一个笼统要求，而是拆分到不同任务中，让每个生成环节都有清楚的角色、边界和输出格式。",
            False,
        ),
        (
            "3.3 反馈闭环：从模型输出到可调试流程",
            "响应反馈机制体现为四个层次",
            "单次生成并不能保证稳定，系统需要在格式、内容、产物和用户操作四个层面形成反馈，这样才能把 AI 输出纳入可调试的工程流程。",
            False,
        ),
        (
            "3.4 智能性边界：可控 Agent 的具体体现",
            "从系统思路看，CourseAgent 的智能性主要体现在三个方面",
            "在本项目中，智能性并不是无限自主，而是在课程资料处理这个明确场景中，系统能够理解目标、选择工具、生成产物并记录结果。",
            False,
        ),
        (
            "4.1 工作台划分：文档生成与视频分析",
            "核心功能围绕两个工作台展开",
            "前文说明了系统架构和提示词策略，本节转向用户能直接看到的界面。UI 设计的重点是把复杂流程压缩成清楚的工作台入口。",
            False,
        ),
        (
            "4.2 文档生成界面：从上传到产物概览",
            "图 2 课程文档生成与产物概览页面",
            "文档生成是 CourseAgent 最核心的入口之一，界面需要同时体现任务目标、运行环境、生成进度和最终产物，帮助用户确认一次生成是否完整。",
            True,
        ),
        (
            "4.3 视频分析界面：从课堂录屏到知识复盘",
            "图 4 课堂视频分析工作台",
            "课堂视频分析关注的是把长视频变成可复习的知识结构，因此 UI 不只提供上传入口，还要突出字幕、时间戳、导图和 highlight 的关系。",
            True,
        ),
        (
            "4.4 产物与运行诊断：保证结果可回溯",
            "图 7 产物中心与历史批次管理",
            "当系统产物变多后，历史记录和运行详情就不只是辅助页面，而是保证课堂展示版本稳定、错误可定位、结果可复用的重要机制。",
            True,
        ),
        (
            "5.1 验证目标：从功能可用到质量可解释",
            "系统验证围绕功能完整性、生成质量、稳定性和交互可用性展开",
            "实验部分不只检查文件是否生成，还要观察生成结果是否贴合资料、系统在异常情况下是否能解释原因，以及用户是否能顺利完成完整流程。",
            False,
        ),
        (
            "5.2 实验现象：提示词与工程闭环共同影响质量",
            "从实验现象看，CourseAgent 的效果提升主要来自两个方面",
            "通过测试可以看到，模型能力只是一个因素，提示词边界、结构化输出、文件校验和错误记录同样会影响最终产物质量。",
            False,
        ),
        (
            "5.3 阶段结论：课程辅助 Agent 的可行性",
            "总的来说，CourseAgent 已经具备课程辅助 Agent 的基本形态",
            "在完成文档和视频两个主要场景验证后，可以对系统做阶段性判断：它已经能承担课程资料初稿生成和复盘整理任务，但仍需要继续提高质量和效率。",
            False,
        ),
        (
            "6.1 当前不足：质量、性能与架构约束",
            "当前系统的主要问题可以分为三类",
            "最后需要客观看待系统边界。当前版本已经能完成课程项目演示，但在模型理解深度、长任务性能和上线架构方面仍然有明显提升空间。",
            False,
        ),
        (
            "6.2 改进路径：知识增强、质量评估与后台化",
            "后续改进可以从以下方向展开",
            "针对这些不足，改进方向应同时覆盖内容质量、系统性能和产品体验，而不是只继续增加功能按钮。",
            False,
        ),
        (
            "6.3 后续演进：从本地项目到可上线平台",
            "从上线架构看，需要把本地文件式历史记录升级为数据库驱动的任务系统",
            "如果继续发展为真实可用的平台，CourseAgent 需要从单机工作台演进为具备用户体系、任务队列、对象存储和监控能力的在线系统。",
            False,
        ),
    ]
    for title, anchor, note, before_drawing in additions:
        add_once(doc, title, anchor, note, before_drawing)

    # Directly enforce black headings on all existing and inserted headings.
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            for run in paragraph.runs:
                style_run(run, 14 if paragraph.style.name == "Heading 1" else 12, True, "000000")

    doc.save(docx)
    print(docx)


if __name__ == "__main__":
    main()
