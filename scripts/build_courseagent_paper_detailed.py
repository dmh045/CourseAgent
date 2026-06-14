from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from build_courseagent_paper import (
    DIAGRAM_PATH,
    add_bullet,
    add_center,
    add_paragraph,
    configure_styles,
    create_diagram,
    set_cell_shading,
    set_cell_text,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
SCREENSHOT_DIR = OUT_DIR / "ui_screenshots"
DOCX_PATH = OUT_DIR / "CourseAgent_课程项目小论文_技术增强版.docx"


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    for line in text.strip().splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_screenshot(doc: Document, filename: str, caption: str) -> None:
    path = SCREENSHOT_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)


def build_document() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    create_diagram()

    doc = Document()
    configure_styles(doc)

    add_center(doc, "基于大语言模型的课程资料智能生成与课堂视频分析系统设计与实现", 20, True, "0B2545")
    add_center(doc, "题目：基于大语言模型的 CourseAgent 课程辅助系统", 12)
    add_center(doc, "学号：______________    姓名：______________", 12)
    add_center(doc, "课程项目小论文（技术增强版）", 11, False, "475569")
    doc.add_paragraph()

    doc.add_heading("摘要", level=1)
    add_paragraph(
        doc,
        "本文围绕 CourseAgent 课程辅助系统的设计与实现展开，系统面向课程文档整理、课堂展示材料生成和课堂视频复习三个典型学习场景，综合使用大语言模型、文档解析、结构化提示词、PPT 自动生成、MoviePy 视频合成、本地语音识别和运行产物管理等技术。系统前端采用 Streamlit 构建交互式工作台，后端采用模块化工具链与 Agent 工作流组织业务逻辑。文档处理流程通过 Planner、Tool Executor、Verifier、Auto-Repair 和 Report 节点完成任务规划、工具调用、产物校验和结果记录；课堂视频流程通过 SRT 解析或本地 Whisper 识别获得字幕，再进行知识点总结、重点时间戳标记、highlight 片段提取和思维导图生成。系统支持 OpenAI、Gemini、DeepSeek 等兼容接口，并在 API 不可用时提供本地规则降级，同时将每次运行保存到独立目录，避免演示产物互相覆盖。实验和调试结果表明，CourseAgent 能够较完整地完成课程资料到展示材料的自动化转化，但在长视频耗时、模型输出深度、PPT 设计美观度和多用户任务调度方面仍有改进空间。",
    )
    add_paragraph(doc, "关键词：CourseAgent；大语言模型；Agent 工作流；提示词工程；课堂视频分析；PPT 自动生成；Streamlit")

    doc.add_heading("1. 引言", level=1)
    add_paragraph(
        doc,
        "在课程学习和课堂展示中，学生通常需要从 PDF 课件、课程讲义、课堂录屏和参考资料中提取知识点，并进一步整理成摘要、思维导图、展示 PPT、答辩讲稿和复习提纲。这一过程具有明显的重复性和流程性，但又要求输出内容紧贴原始资料，不能脱离课程主题自由发挥。传统做法主要依赖人工阅读和手工编辑，效率较低，并且容易出现知识结构不清、PPT 内容空泛、讲稿与幻灯片重复、课堂视频难以定位重点等问题。",
    )
    add_paragraph(
        doc,
        "CourseAgent 的设计目标是把课程资料处理流程抽象为可执行的智能工作流。与简单的聊天式问答不同，本系统需要完成的是一组可落地的文件产物，包括 Markdown 摘要、关键词 JSON、思维导图图片、PPTX、讲稿、MP4 讲解视频、SRT 字幕、时间戳 JSON、运行报告和历史记录 manifest。因此，系统必须同时解决自然语言理解、结构化输出、文件生成、产物校验、异常恢复和 UI 交互等问题。",
    )
    add_paragraph(
        doc,
        "本文按照系统开发过程展开，首先介绍需求和技术架构，然后分析 Agent 工作流、提示词设计、文档生成链路和视频分析链路，最后结合 UI 演示截图说明系统交互，并总结实验验证结果、系统问题和后续改进方向。",
    )

    doc.add_heading("2. 需求分析与系统目标", level=1)
    add_paragraph(
        doc,
        "系统需求可以分为功能性需求和非功能性需求。功能性需求强调用户能否一键完成学习材料生成；非功能性需求强调结果是否可追溯、运行是否稳定、失败是否可恢复，以及课堂展示前是否能保存确定版本。",
    )
    add_table(
        doc,
        ["需求类型", "具体需求", "系统实现"],
        [
            ("输入处理", "支持课程文档和课堂视频两类输入。", "文档支持 TXT、DOCX、PDF；视频支持 MP4、MOV、MKV，并允许上传 SRT 字幕。"),
            ("内容生成", "从课程资料生成摘要、关键词、思维导图、PPT 和讲稿。", "由 summary_tool、keyword_tool、mindmap_tool、ppt_outline_tool、script_tool、ppt_tool 等模块完成。"),
            ("视频复习", "对课堂录屏识别字幕并提取知识点和重点时间戳。", "video_analysis_tool 负责字幕转录、字幕转文本、LLM 分析、markers 和 highlights 输出。"),
            ("产物管理", "避免多次分析互相覆盖，便于课堂演示。", "history_tool 为每次运行创建独立 run_dir，并写入 manifest.json。"),
            ("异常处理", "API 或本地模型失败时系统不应整体崩溃。", "LLM 调用提供 fallback；语音识别使用子进程隔离和 PyTorch Whisper 备用后端。"),
            ("交互体验", "用户应能从 UI 页面一键完成工作流。", "Streamlit 页面提供文档工作台、视频分析、产物中心和运行详情四个主 tab。"),
        ],
        [1600, 3550, 4210],
    )

    doc.add_heading("3. 系统总体架构", level=1)
    add_paragraph(
        doc,
        "CourseAgent 采用分层架构。最上层是 Streamlit UI，负责文件上传、任务目标输入、API 配置、运行状态展示、产物下载和历史记录管理。中间层是 Agent 工作流，由 agent_graph.py 维护状态并按照节点顺序执行。底层是工具模块，分别负责文档解析、内容生成、PPT 生成、视频合成、字幕识别、思维导图渲染和报告保存。配置层 config.py 统一管理运行目录、上传目录、输出目录、缓存目录和 API 供应商预设。",
    )
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "图 1  CourseAgent 系统总体设计框架")
    add_paragraph(
        doc,
        "系统运行目录被统一迁移到 E:\\CourseAgent，其中 uploads 保存上传文件，output 保存运行产物，output\\runs 按批次保存历史记录，cache 保存模型缓存、语音识别缓存和临时文件。这种设计把项目源码目录和大文件运行目录分离，既减少 C 盘压力，也降低测试产物污染源码目录的风险。",
    )
    add_code_block(
        doc,
        """
RUNTIME_DIR = E:\\CourseAgent
UPLOAD_DIR  = E:\\CourseAgent\\uploads
OUTPUT_DIR  = E:\\CourseAgent\\output
CACHE_DIR   = E:\\CourseAgent\\cache
RUNS_DIR    = E:\\CourseAgent\\output\\runs
        """,
    )

    doc.add_heading("4. Agent 工作流设计", level=1)
    add_paragraph(
        doc,
        "系统的核心控制逻辑位于 agent_graph.py。它定义 AgentState 作为跨节点共享状态，记录用户目标、文件路径、文档文本、动态计划、已选工具、各类生成结果、校验结果、修复记录、运行日志、trace、错误列表和运行耗时。每个节点接收并返回同一个 state，从而形成可追踪的流水线。",
    )
    add_paragraph(
        doc,
        "工作流包括五类关键节点：read_document_node 读取并预览原文；planner_node 根据用户目标和文档预览生成工具计划；executor_node 根据计划顺序调用工具；verifier_node 检查计划内产物是否真实存在并非空；final_node 写入运行报告和 manifest。该流程体现了 Agent 的基本特征：不是固定按钮触发单个函数，而是根据目标动态选择工具，并对中间产物进行验证和修复。",
    )
    add_table(
        doc,
        ["节点", "输入", "处理逻辑", "输出"],
        [
            ("Read", "上传文件路径", "根据后缀调用 TXT、DOCX 或 PDF 解析函数，生成 document_text 和 preview。", "文档全文与预览"),
            ("Planner", "用户目标、文档预览", "调用 plan_task，结合显式目标补齐必要工具依赖。", "steps、selected_tools"),
            ("Executor", "工具计划", "通过 TOOL_REGISTRY 顺序执行摘要、关键词、导图、PPT、视频等工具。", "各类中间结果和文件"),
            ("Verifier", "计划工具和产物路径", "调用 verify_outputs 检查文件存在性和大小，缺失时触发 _auto_repair。", "verification、repairs"),
            ("Report", "最终 state", "汇总 logs、trace、errors、artifacts，写入 run_report.md 和 manifest.json。", "运行报告和历史记录"),
        ],
        [1200, 1900, 4300, 1960],
    )
    add_paragraph(
        doc,
        "动态规划还处理了下游依赖。例如用户只要求生成视频时，系统会自动补齐 PPT 大纲和逐页讲稿，因为视频生成依赖这些中间产物；用户只要求思维导图时，系统不会额外生成 PPT 和视频，从而减少无关耗时。这一点在测试用例 test_dynamic_goal_mindmap_only 和 test_dynamic_goal_video_adds_dependencies 中得到验证。",
    )

    doc.add_heading("5. LLM 接口适配与提示词工程", level=1)
    add_paragraph(
        doc,
        "LLM 调用集中封装在 tools/llm_client.py。系统支持 OpenAI 及兼容 OpenAI 协议的供应商，通过 provider、base_url、model、api_type 和 api_key 统一配置。接口层优先尝试 Responses API，并支持结构化 JSON schema；当供应商不支持 response_format 或 Responses API 时，系统自动退回 Chat Completions 路径。若 API Key 缺失或调用失败，则返回本地 fallback，保证工作流不中断。",
    )
    add_paragraph(
        doc,
        "为了避免模型输出不可解析，系统为 planner、summary、keywords、mindmap、ppt_outline、slide_scripts 等任务定义了 JSON_SCHEMAS。提示词不仅要求模型输出内容，还明确约束字段结构、数组元素、标题、说明和证据等字段。这样做的目的不是追求模型回答的自由度，而是保证后续工具可以稳定消费模型结果。",
    )
    add_bullet(doc, "文档摘要提示词：要求严格依据上传资料，避免把系统功能、Agent 工作流或无关背景写入课程内容。")
    add_bullet(doc, "PPT 大纲提示词：要求生成适合课堂答辩的页面结构，每页包含标题和可讲解要点。")
    add_bullet(doc, "讲稿提示词：强调使用自然教学语言解释 PPT 内容，而不是逐字复述幻灯片。")
    add_bullet(doc, "思维导图提示词：要求层级化组织课程主题、核心概念、方法步骤、示例和结论。")
    add_bullet(doc, "视频分析提示词：要求从字幕中提取 key_points、markers 和 highlights，并禁止编造字幕外内容。")
    add_paragraph(
        doc,
        "这种提示词设计使 CourseAgent 的智能性体现在“任务约束与工具协同”上：模型负责理解和组织内容，程序负责解析、生成文件和校验结果，两者结合后形成可重复执行的课程资料生产流程。",
    )

    doc.add_heading("6. 文档处理链路实现", level=1)
    add_paragraph(
        doc,
        "文档处理从 tools/document_tool.py 开始。TXT 文件会依次尝试 utf-8-sig、utf-8 和 gbk 编码；DOCX 文件通过 python-docx 读取段落和表格；PDF 文件通过 PyMuPDF 获取页面文本。解析后的文本先保存为 document_text，再截取前 1500 个字符作为 document_preview 给 Planner 使用，避免规划阶段消耗过长上下文。",
    )
    add_paragraph(
        doc,
        "内容生成链路依次产生 summary.md、keywords.json、mindmap.json、mindmap.png、ppt_outline.json、speech_script.md、generated_presentation.pptx、final_video.mp4、subtitles.srt、video_markers.json、video_chapters.md、voice_report.json 和 run_report.md。PPT 生成模块将大纲、讲稿和思维导图合并为正式演示文稿；视频生成模块使用 MoviePy 将每页内容渲染为动态帧，并可结合语音文件生成 MP4，同时输出字幕和章节时间戳。",
    )
    add_paragraph(
        doc,
        "为了降低重复运行成本，系统对文档任务计算 cache_fingerprint。指纹由文件内容、用户目标、供应商、模型和 mock_mode 共同决定。如果同一文档和同一目标之前已成功生成，系统可以复用历史成功产物到新的 run_dir 中，从而显著降低课堂展示前反复调试的等待时间。",
    )

    doc.add_heading("7. 课堂视频分析链路实现", level=1)
    add_paragraph(
        doc,
        "课堂视频分析由 tools/video_analysis_tool.py 实现。若用户上传 SRT，系统直接解析字幕；若未上传 SRT 且启用本地识别，则调用 transcribe_video_to_subtitles 生成 course_video_auto_subtitles.srt。随后系统将字幕转换为带时间范围的 transcript_text，进一步生成 course_video_transcript.txt、course_video_analysis.md、course_video_markers.json、course_video_highlights.md、course_video_mindmap.json 和 course_video_mindmap.png。",
    )
    add_paragraph(
        doc,
        "本地语音识别采用隔离子进程设计。最初使用 faster-whisper/ctranslate2，但在 Windows + Anaconda 环境下可能发生 MSVCP140.dll 级别的原生崩溃，这类错误无法被普通 Python try/except 捕获。因此系统将识别过程放入 whisper_worker.py 子进程，即使底层崩溃也不会杀死 Streamlit 主进程。之后又增加 torch_whisper_worker.py 作为 PyTorch Whisper 备用后端，并通过 COURSEAGENT_TRANSCRIBE_BACKEND 控制默认后端选择。",
    )
    doc.add_page_break()
    doc.add_heading("7.1 视频分析模块关键环节", level=2)
    add_table(
        doc,
        ["环节", "实现方式", "容错策略"],
        [
            ("字幕来源", "优先使用用户上传 SRT；缺失时本地识别。", "用户可绕过本地识别，直接上传字幕保证稳定。"),
            ("识别缓存", "按视频文件哈希、模型大小和语言生成 transcript 缓存。", "同一视频再次分析可复用字幕，减少耗时。"),
            ("识别进程", "faster-whisper 和 PyTorch Whisper 均以独立 worker 运行。", "子进程崩溃不影响主 UI。"),
            ("知识点分析", "字幕转 transcript 后调用 LLM 或本地规则分析。", "无 API 时使用本地分段、词频和时间块策略生成基础结果。"),
            ("结果表达", "生成 markers、highlights、summary 和 mindmap。", "UI 可下载 JSON、SRT 与分析报告。"),
        ],
        [1500, 4300, 3560],
    )

    doc.add_page_break()
    doc.add_heading("8. UI 页面与交互演示", level=1)
    add_paragraph(
        doc,
        "系统 UI 分为左侧运行控制区和右侧主工作区。左侧展示当前供应商、模型、运行模式、API 配置和工作流说明；右侧顶部展示运行环境、API 状态、输出目录和系统简介。主工作区通过四个 tab 组织主要功能：课程文档生成、课堂视频分析、产物中心和运行详情。",
    )
    add_screenshot(doc, "ui_01_document_workspace.png", "图 2  课程文档生成工作台：上传课程资料、填写任务目标并一键生成全部材料")
    add_paragraph(
        doc,
        "文档生成页面把流程拆成上传资料、Agent 规划、内容生成、产物校验和下载交付五个步骤，使用户能够理解系统不是直接生成一个文件，而是执行一条完整工具链。任务目标文本框允许用户明确 PPT 时长、输出材料和讲解风格，从而影响 Planner 的工具选择。",
    )
    add_screenshot(doc, "ui_02_video_analysis.png", "图 3  课堂视频分析工作台：上传录屏或字幕，生成知识点、导图、时间戳和重点标记")
    add_paragraph(
        doc,
        "视频分析页面突出 SRT 可选上传和本地识别模型选择。为了减少等待和失败风险，系统建议优先上传字幕；当用户没有字幕时，再调用本地识别后端。识别完成后，系统会在页面中展示视频、知识点摘要、markers 和 highlights，并提供分析报告下载。",
    )
    add_screenshot(doc, "ui_03_asset_history.png", "图 4  产物中心：按运行批次保存历史记录，支持课堂展示版本固定与下载")
    add_screenshot(doc, "ui_04_run_details.png", "图 5  运行详情：展示 Agent 决策、执行日志、校验结果和错误信息")
    add_paragraph(
        doc,
        "产物中心的设计解决了课堂展示前最实际的问题：提前生成的文档产物和视频分析产物不会互相覆盖。每次运行都会形成独立 run_id，manifest 中记录类型、状态、源文件、目标、产物路径、耗时、错误和校验结果。运行详情页则用于排查质量问题或 API 问题，适合查看 Planner 决策是否合理、工具是否执行成功、Verifier 是否发现缺失产物。",
    )

    doc.add_heading("9. 实验验证与结果分析", level=1)
    add_paragraph(
        doc,
        "系统验证主要从单元测试、集成运行和异常调试三个层面展开。单元测试覆盖文档读取、LLM fallback、动态计划补齐、思维导图生成、PPT 生成、视频分析和产物校验等模块。最近一次显式运行 python -m pytest tests -q 的结果为 18 passed，说明核心模块在测试输入下能够正常工作。",
    )
    add_table(
        doc,
        ["验证对象", "验证内容", "结果"],
        [
            ("动态规划", "只要求思维导图时不生成 PPT/视频；要求视频时自动补齐 PPT 大纲和讲稿。", "通过测试，说明 Planner 与依赖补齐逻辑有效。"),
            ("文档产物", "检查摘要、关键词、思维导图、PPT、讲稿、视频和报告是否生成。", "Verifier 能发现缺失文件并触发修复。"),
            ("视频分析", "上传 SRT 后生成转录、分析报告、markers、highlights 和导图。", "测试通过，说明字幕分析链路可用。"),
            ("历史记录", "每次运行写入 manifest 并在产物中心展示。", "可按 run_id 区分任务，避免产物覆盖。"),
            ("异常恢复", "API 失败、本地识别崩溃、缺失产物等异常。", "主工作台保持运行，错误写入日志或 manifest。"),
        ],
        [1700, 4700, 2960],
    )
    add_paragraph(
        doc,
        "在实际调试中，系统曾遇到 Gemini 额度限制、DeepSeek response_format 不兼容、executor_node 缺少 json 导入、Streamlit DuplicateElementKey、视频识别导致主进程断联等问题。通过接口降级、统一 JSON 解析、唯一 key、子进程隔离、PyTorch Whisper 备用后端和历史记录机制，这些问题逐步被定位和修复。该过程说明，对于 AI 应用而言，工程可靠性与模型能力同样重要。",
    )

    doc.add_heading("10. 系统存在问题与改进方向", level=1)
    add_paragraph(
        doc,
        "当前系统已经具备课程辅助 Agent 的基本形态，但仍存在若干不足。第一，生成质量高度依赖模型能力和 API 额度，在弱模型或本地规则模式下，内容可能偏概括，难以达到高质量论文或课堂展示标准。第二，PPT 自动排版虽然能生成可用文件，但视觉设计仍与人工精修存在差距。第三，视频识别和视频合成耗时较长，尤其是长视频首次识别需要等待模型加载和音频转录。第四，系统目前主要面向单用户本地运行，尚未提供真正的后台任务队列、多用户隔离和权限管理。",
    )
    add_paragraph(
        doc,
        "后续改进可以从五个方向展开：一是引入向量检索和分段总结，对长文档建立可引用的知识片段索引；二是加入多轮自我审查，让模型对 PPT 大纲、讲稿和思维导图进行事实一致性复核；三是建立 PPT 模板系统，根据课程类型自动选择视觉布局和图表；四是引入异步任务队列，将耗时的视频识别和合成放到后台执行，并在 UI 中显示更细粒度进度；五是进一步优化本地语音识别环境，提供 faster-whisper、PyTorch Whisper 和外部字幕 API 的多后端切换。",
    )
    add_paragraph(
        doc,
        "总体来看，CourseAgent 的价值不只是把几个生成函数嵌入页面，而是把大语言模型、文件处理、视频处理、校验修复和历史管理组织成一个面向课程场景的完整智能工作流。它能够把课程资料从“原始输入”转化为“可展示、可讲解、可复习、可追踪”的学习产物，体现了 AI 在教育辅助工具中的实际应用潜力。",
    )

    doc.add_heading("参考资料", level=1)
    add_paragraph(doc, "[1] OpenAI. Large Language Models and Prompt Engineering Practices.")
    add_paragraph(doc, "[2] Streamlit Documentation. Building interactive data and AI applications.")
    add_paragraph(doc, "[3] Whisper Speech Recognition Model and Related Open Source Implementations.")
    add_paragraph(doc, "[4] MoviePy Documentation. Video editing and compositing in Python.")
    add_paragraph(doc, "[5] CourseAgent 项目源码：app.py、agent_graph.py、config.py、tools 模块与 tests 测试用例。")
    add_paragraph(doc, "[6] 软件工程与交互设计课程资料：KLM 效率模型、交互设计过程等。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CourseAgent 课程项目小论文（技术增强版）")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
