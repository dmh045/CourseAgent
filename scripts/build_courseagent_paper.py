from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "CourseAgent_课程项目小论文.docx"
DIAGRAM_PATH = OUT_DIR / "courseagent_architecture.png"


TITLE = "基于大语言模型的课程资料智能生成与课堂视频分析系统设计与实现"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, body: str, fill: str) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline="#9FB6D8", width=2)
    draw.text((x1 + 22, y1 + 18), title, font=font(24, True), fill="#0B2545")
    body_font = font(18)
    y = y1 + 56
    for line in wrap_text(draw, body, body_font, x2 - x1 - 44):
        draw.text((x1 + 22, y), line, font=body_font, fill="#334155")
        y += 26


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#2563EB", width=4)
    ex, ey = end
    sx, sy = start
    if ex > sx:
        points = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)]
    elif ex < sx:
        points = [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
    elif ey > sy:
        points = [(ex, ey), (ex - 8, ey - 14), (ex + 8, ey - 14)]
    else:
        points = [(ex, ey), (ex - 8, ey + 14), (ex + 8, ey + 14)]
    draw.polygon(points, fill="#2563EB")


def create_diagram() -> None:
    img = Image.new("RGB", (1800, 1080), "#F6F8FC")
    draw = ImageDraw.Draw(img)
    draw.text((70, 48), "CourseAgent 系统总体设计框架", font=font(42, True), fill="#0B2545")
    draw.text((70, 110), "从课程文档与课堂视频输入出发，经 Agent 规划、工具执行、校验修复，最终形成可展示、可复习、可追溯的学习产物。", font=font(24), fill="#475569")

    boxes = {
        "input": (70, 210, 410, 390),
        "planner": (520, 210, 860, 390),
        "tools": (970, 210, 1320, 390),
        "verify": (1410, 210, 1730, 390),
        "doc": (210, 560, 540, 790),
        "video": (730, 560, 1060, 790),
        "assets": (1250, 560, 1580, 790),
    }
    draw_box(draw, boxes["input"], "输入层", "课程 PDF / DOCX / TXT\n课堂录屏 MP4 / SRT 字幕\n任务目标与 API 配置", "#FFFFFF")
    draw_box(draw, boxes["planner"], "Agent 规划层", "理解用户目标\n拆解任务步骤\n选择文档、PPT、视频等工具", "#EEF6FF")
    draw_box(draw, boxes["tools"], "工具执行层", "文档解析、LLM 生成\n思维导图、PPT、讲稿\n字幕识别与时间戳分析", "#ECFDF5")
    draw_box(draw, boxes["verify"], "校验修复层", "检查产物完整性\n发现缺失自动修复\n输出运行报告", "#FFF7ED")
    draw_box(draw, boxes["doc"], "文档工作台", "摘要、关键词、课程大纲\n答辩 PPT、逐页讲稿\n讲解视频与字幕文件", "#FFFFFF")
    draw_box(draw, boxes["video"], "视频工作台", "语音识别生成字幕\n知识点总结与思维导图\n重点时间戳与 Highlight 标记", "#FFFFFF")
    draw_box(draw, boxes["assets"], "历史记录中心", "按运行批次保存产物\n支持固定展示版本\n便于课堂演示与回溯", "#FFFFFF")

    arrow(draw, (410, 300), (520, 300))
    arrow(draw, (860, 300), (970, 300))
    arrow(draw, (1320, 300), (1410, 300))
    arrow(draw, (1135, 390), (375, 560))
    arrow(draw, (1135, 390), (895, 560))
    arrow(draw, (1570, 390), (1415, 560))

    draw.rounded_rectangle((70, 880, 1730, 1010), radius=22, fill="#E8EEF5", outline="#CBD5E1", width=2)
    draw.text((105, 910), "核心思想", font=font(26, True), fill="#0B2545")
    draw.text(
        (245, 912),
        "将大语言模型能力封装进可控工作流：提示词约束内容质量，工具模块负责具体产物，校验节点保证结果完整，历史记录支撑展示与复用。",
        font=font(24),
        fill="#334155",
    )
    img.save(DIAGRAM_PATH, quality=95)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    if bold:
        run.font.color.rgb = RGBColor(11, 37, 69)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_paragraph(doc: Document, text: str, style: str | None = None, bold_lead: str | None = None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None else None
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        p.add_run(text[len(bold_lead) :])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_center(doc: Document, text: str, size: int, bold: bool = False, color: str = "000000") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor.from_string(color)


def build_document() -> None:
    create_diagram()
    doc = Document()
    configure_styles(doc)

    add_center(doc, TITLE, 20, True, "0B2545")
    add_center(doc, "题目：基于大语言模型的 CourseAgent 课程辅助系统", 12)
    add_center(doc, "学号：______________    姓名：______________", 12)
    add_center(doc, "课程项目小论文", 11, False, "475569")
    doc.add_paragraph()

    doc.add_heading("摘要", level=1)
    add_paragraph(
        doc,
        "随着大语言模型、语音识别和自动化内容生成技术的发展，课程资料的整理方式正在从人工编辑逐步转向智能辅助。本文围绕 CourseAgent 系统的设计与实现展开，介绍一个面向课程文档与课堂视频的智能工作台。系统支持上传课程 PDF、DOCX、TXT 文档，自动生成摘要、关键词、思维导图、答辩 PPT、逐页讲稿、讲解视频、字幕、重点时间戳和运行报告；同时支持上传课堂录屏，对视频进行字幕识别、知识点总结、重点片段标记和复习结构生成。系统采用 Streamlit 构建交互式界面，将大语言模型能力封装为可控的 Agent 工作流，并通过任务规划、工具执行、产物校验、失败修复和历史记录管理提升系统可用性。实验表明，该系统能够有效降低课程资料整理与展示材料制作的时间成本，但在长视频处理、模型稳定性、PPT 美观度和深度知识分析方面仍有进一步优化空间。",
    )
    add_paragraph(doc, "关键词：CourseAgent；大语言模型；课程资料生成；课堂视频分析；提示词工程；智能工作流")

    doc.add_heading("1. 引言", level=1)
    add_paragraph(
        doc,
        "在日常课程学习和课堂展示中，学生往往需要花费大量时间完成资料阅读、知识点提炼、PPT 制作、讲稿撰写和课堂视频复习等工作。尤其是在软件工程、人机交互、系统设计等课程中，资料内容通常具有概念多、结构复杂、示例分散的特点，单纯依靠人工整理不仅效率较低，也容易出现知识点遗漏、表达不清和展示材料不统一等问题。",
    )
    add_paragraph(
        doc,
        "CourseAgent 的开发背景正是来源于这一学习场景。系统希望将大语言模型的文本理解能力、文档解析能力、PPT 自动生成能力、视频处理能力和语音识别能力结合起来，构建一个面向课程资料处理的智能工作台。用户只需要上传课程文档或课堂录屏，并输入任务目标，系统即可自动规划处理步骤，生成可用于课堂展示、期末复习或学习归档的一组产物。",
    )
    add_paragraph(
        doc,
        "从应用意义上看，该系统不仅可以提高学习材料整理效率，还能够帮助用户形成更清晰的知识结构。对于课堂展示而言，系统生成的 PPT、讲稿和时间戳可以减少准备成本；对于课后复习而言，摘要、关键词、思维导图和视频 highlight 标记可以帮助用户快速定位重点内容。与单次调用聊天模型不同，CourseAgent 更强调工作流集成，即把多个 AI 与工具模块组织成一个可执行、可校验、可追踪的 Agent 系统。",
    )

    doc.add_heading("2. AI + 系统设计", level=1)
    add_paragraph(
        doc,
        "CourseAgent 采用前后端一体化的轻量级架构。前端使用 Streamlit 构建 Web UI，负责文件上传、任务输入、API 配置、产物预览、历史记录展示和下载交互；后端由多个工具模块组成，包括文档解析、摘要生成、关键词提取、思维导图生成、PPT 生成、讲稿生成、讲解视频生成、字幕处理、课堂视频分析和历史记录管理等。",
    )
    add_paragraph(
        doc,
        "系统的核心并不是简单地把各个功能按钮堆叠在一起，而是引入 Agent 工作流。整体流程可以概括为 Planner → Tool Executor → Verifier → Auto-Repair → Report。Planner 根据用户上传内容和任务目标规划要执行的步骤；Tool Executor 调用具体工具模块生成产物；Verifier 检查摘要、关键词、PPT、思维导图、视频、字幕和报告是否完整；Auto-Repair 在发现缺失或异常时尝试补齐；Report 最终保存运行详情和产物清单。",
    )

    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.3))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph("图 1  CourseAgent 系统总体设计框架")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    add_paragraph(
        doc,
        "在文档处理路径中，系统首先读取上传资料并生成文本预览，然后调用大语言模型或本地规则生成摘要、关键词、PPT 大纲和逐页讲稿，最后通过 PPT 与媒体生成工具输出可下载文件。在视频分析路径中，系统读取用户上传的课堂视频，如果用户没有提供 SRT 字幕，则调用本地语音识别后端生成字幕，再基于字幕文本提取知识点、时间戳、重点片段和思维导图。",
    )
    add_paragraph(
        doc,
        "为了适应不同 API 条件，系统支持 OpenAI、Gemini、DeepSeek 等接口配置，同时保留本地提取式规则作为降级方案。当 API Key 缺失、额度不足或连接失败时，系统仍能生成基础摘要和部分结构化产物。为了适应课堂展示需要，系统还引入按运行批次保存的历史记录机制，避免不同任务产物相互覆盖。",
    )

    doc.add_heading("3. AI 提示词设计与响应反馈", level=1)
    add_paragraph(
        doc,
        "提示词设计是 CourseAgent 智能性的关键。系统并不是让模型自由发挥，而是通过分任务提示词约束模型的输入依据、输出格式、内容深度和表达风格。文档理解提示词要求模型严格基于原文进行总结，避免将系统自身功能或无关背景写入课程内容；PPT 生成提示词要求每一页具有明确主题、核心论点、支撑解释和讲解逻辑；讲稿提示词要求使用更自然的教学语言，而不是机械朗读幻灯片文字。",
    )
    add_paragraph(
        doc,
        "在思维导图生成中，提示词强调层级结构，要求区分课程主题、核心概念、方法步骤、案例说明和结论，避免只生成几个空泛关键词。在课堂视频分析中，提示词要求模型根据字幕转录内容提取知识点、重点时间戳、片段摘要和 highlight 标记，并明确要求不得编造字幕中不存在的内容。",
    )
    add_paragraph(
        doc,
        "系统还设计了响应反馈机制。首先，模型输出尽量采用 JSON 或结构化文本，便于程序解析和后续工具调用；其次，系统会对关键产物进行完整性校验，例如检查 summary、keywords、mindmap、PPT、讲稿、视频、字幕和运行报告是否存在；最后，如果发现缺失产物，系统会进入自动修复或失败记录流程。这样可以减少一次生成失败导致整个工作流不可用的情况。",
    )
    add_paragraph(
        doc,
        "从设计思路上看，CourseAgent 的提示词不是单点提示，而是一组围绕任务链路协同工作的提示词。它们共同服务于同一个目标：把课程资料转化为可展示、可讲解、可复习、可追踪的学习材料。",
    )

    doc.add_heading("4. 核心功能设计", level=1)
    add_paragraph(
        doc,
        "系统核心功能围绕两类输入展开：课程文档和课堂视频。课程文档工作台主要解决“从资料到展示材料”的问题；课堂视频工作台主要解决“从录屏到复习结构”的问题。两类工作流共享 API 配置、历史记录、产物校验和下载管理等基础能力。",
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [1900, 3850, 3610]
    set_table_geometry(table, widths)
    headers = ["功能模块", "主要作用", "输出产物"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    rows = [
        ("文档解析", "读取 PDF、DOCX、TXT 等课程资料，并转化为可供模型理解的文本。", "文本预览、原文片段"),
        ("内容生成", "根据课程资料生成摘要、关键词、PPT 大纲、逐页讲稿和讲解说明。", "摘要、关键词、讲稿、PPT 大纲"),
        ("可视化产物", "将课程结构转化为思维导图和正式 PPT，辅助课堂展示。", "思维导图、PPTX 文件"),
        ("视频讲解", "根据 PPT 和讲稿合成讲解视频，并生成字幕和重点时间戳。", "MP4 视频、SRT 字幕、时间戳"),
        ("课堂视频分析", "对上传的课堂录屏进行语音识别和知识点分析，提取重点片段。", "字幕、知识点总结、highlight 标记"),
        ("历史记录", "按运行批次保存产物，支持展示版本固定和后续回溯。", "运行清单、下载入口、错误记录"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    add_paragraph(
        doc,
        "在实现层面，系统采用模块化设计，每个功能模块负责一个相对独立的任务。例如 document_tool 负责文档读取，summary_tool 和 keyword_tool 负责内容提炼，ppt_tool 负责 PPT 生成，mindmap_tool 负责知识结构可视化，video_analysis_tool 负责课堂视频字幕和知识点分析，history_tool 负责运行记录管理。模块化设计使系统便于扩展，也便于针对单个功能进行调试和优化。",
    )
    add_paragraph(
        doc,
        "对于视频分析模块，系统特别考虑了本地语音识别的稳定性问题。由于部分 Windows 环境下 faster-whisper/ctranslate2 可能发生底层 DLL 崩溃，系统将语音识别放入独立子进程，并增加 PyTorch Whisper 备用后端。这样即使识别模块失败，也不会直接导致主工作台崩溃，提升了系统整体鲁棒性。",
    )

    doc.add_heading("5. 实验、验证或数据分析与结论", level=1)
    add_paragraph(
        doc,
        "为了验证系统可用性，选取软件工程与交互设计相关课程资料作为测试对象，包括 KLM 效率模型、交互设计过程等 PDF 文档，以及课堂录屏视频。实验主要关注四个方面：产物完整性、内容相关性、交互可用性和运行稳定性。",
    )
    add_bullet(doc, "产物完整性：系统能够按照任务规划生成摘要、关键词、思维导图、PPT、讲稿、字幕、时间戳和运行报告等文件。")
    add_bullet(doc, "内容相关性：在接入大语言模型 API 后，生成内容能够更好地围绕上传课程资料展开；在无 API 情况下，本地规则模式仍可生成基础结果，但深度和表达质量有限。")
    add_bullet(doc, "交互可用性：用户可以在同一页面完成文件上传、任务输入、API 配置、结果预览和下载，历史记录中心也降低了课堂展示前产物被覆盖的风险。")
    add_bullet(doc, "运行稳定性：系统通过子进程隔离、本地缓存、失败记录和备用识别后端缓解了视频处理时间长、模型崩溃和网络不稳定等问题。")
    add_paragraph(
        doc,
        "实验结果表明，CourseAgent 能够显著减少课程资料整理和展示材料制作的重复劳动。对于普通课程 PDF，系统可以在较短时间内生成初步可用的摘要、PPT 和讲稿；对于课堂录屏，系统能够通过字幕识别和知识点提取形成复习线索。总体而言，该系统已经具备课程辅助 Agent 的基本形态，即能够理解任务、规划步骤、调用工具、检查产物并保留运行记录。",
    )
    add_paragraph(
        doc,
        "同时也应看到，系统生成质量仍然受到输入资料质量、模型能力、API 额度、提示词设计和本地运行环境的影响。特别是在 PPT 视觉美观度、长文档深度理解、长视频处理速度和模型输出一致性方面，仍需要进一步优化。",
    )

    doc.add_heading("6. 系统存在问题与改进", level=1)
    add_paragraph(
        doc,
        "当前系统仍存在一些不足。第一，生成内容在无 API 或弱模型条件下容易偏概括，难以达到高质量课堂展示材料的要求。第二，视频识别和视频合成耗时较长，对本地环境依赖较强。第三，PPT 排版虽然可以自动生成，但与人工精心设计的展示文稿相比，在视觉风格、图文布局和重点突出方面仍有差距。第四，系统目前主要通过文件级历史记录保存产物，未来还可以进一步支持版本对比、运行参数复用和多用户管理。",
    )
    add_paragraph(
        doc,
        "针对上述问题，后续可以从以下方向改进：一是引入向量检索和分段阅读机制，提高长文档理解能力；二是加入多轮自我审查与事实一致性校验，减少模型幻觉；三是设计更专业的 PPT 模板系统，提高展示材料美观度；四是引入异步任务队列和进度可视化，减少用户等待焦虑；五是继续优化本地语音识别后端，使课堂视频分析更加稳定。",
    )
    add_paragraph(
        doc,
        "综上，CourseAgent 是一个将大语言模型与课程资料处理场景结合的综合性系统。它的价值不仅在于单个功能的自动化，更在于把文档理解、内容生成、视频分析、产物校验和历史管理组织成完整工作流。随着提示词、模型能力和系统工程能力的进一步提升，该系统可以继续向更智能、更稳定、更适合真实课堂使用的学习辅助 Agent 演进。",
    )

    doc.add_heading("参考资料", level=1)
    add_paragraph(doc, "[1] OpenAI. Large Language Models and Prompt Engineering Practices.")
    add_paragraph(doc, "[2] Streamlit Documentation. Building interactive data and AI applications.")
    add_paragraph(doc, "[3] Whisper Speech Recognition Model and Related Open Source Implementations.")
    add_paragraph(doc, "[4] 软件工程与交互设计课程资料：KLM 效率模型、交互设计过程等。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CourseAgent 课程项目小论文")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
    print(DOCX_PATH)
