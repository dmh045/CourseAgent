from __future__ import annotations

from io import BytesIO
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
TMP_DIR = OUT_DIR / "embedded_ui_figures"


def final_docx_path() -> Path:
    needle = "\u6700\u7ec8\u7248"
    return next(path for path in OUT_DIR.glob("*.docx") if needle in path.name and not path.name.startswith("~$"))


def set_run_font(run, size: float = 10.5, color: str = "1F2937", bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, 9.3, "4B5563")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    r = p.add_run(text)
    set_run_font(r, 10.5, "1F2937")


def add_picture_paragraph(doc: Document, image_path: Path, width: float) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))


def paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def extract_ui_images(doc: Document) -> list[Path]:
    TMP_DIR.mkdir(exist_ok=True)
    paths: list[Path] = []
    # Images 1-2 are architecture / data-flow diagrams. The UI screenshots start at image 3.
    for i, shape in enumerate(doc.inline_shapes, start=1):
        if i < 3:
            continue
        rid = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        part = doc.part.related_parts[rid]
        image = Image.open(BytesIO(part.blob)).convert("RGB")
        out = TMP_DIR / f"ui_embedded_{i:02d}.png"
        image.save(out)
        paths.append(out)
    return paths


def remove_old_ui_block(doc: Document) -> int:
    paragraphs = doc.paragraphs
    start_para = next(p for p in paragraphs if paragraph_text(p).startswith("核心功能围绕两个工作台展开"))
    end_para = next(p for p in paragraphs if paragraph_text(p).startswith("5. 实验"))
    body = doc._body._element
    children = list(body)
    start_idx = children.index(start_para._p) + 1
    end_idx = children.index(end_para._p)
    for element in children[start_idx:end_idx]:
        body.remove(element)
    return start_idx


def insert_elements_at(doc: Document, index: int, elements) -> None:
    body = doc._body._element
    for offset, element in enumerate(elements):
        body.remove(element)
        body.insert(index + offset, element)


def build_new_ui_block(doc: Document, ui_images: list[Path], insert_idx: int) -> None:
    captions = [
        (
            "图 2 课程文档生成与产物概览页面",
            "该图展示文档生成后的工作台状态：页面上方保留运行环境和模型信息，中部呈现文档生成流程，底部集中展示摘要、关键词、PPT、讲稿、视频等产物入口，便于用户确认一次生成是否完整。",
            6.45,
        ),
        (
            "图 3 思维导图、讲解视频与重点内容预览",
            "该图展示生成结果的进一步展开：左侧用思维导图呈现课程知识结构，右侧结合视频预览、时间轴标记和讲解概要，让用户能够从结构和播放两个角度复习资料。",
            5.7,
        ),
        (
            "图 4 课堂视频分析工作台",
            "该图展示课堂视频分析入口。用户可以上传课堂录屏，也可以上传已有 SRT 字幕；系统会根据字幕或本地识别结果继续生成知识点总结、时间戳和 highlight 标记。",
            6.45,
        ),
        (
            "图 5 课堂视频分析结果与知识结构",
            "该图展示视频分析完成后的结果区域：系统将字幕内容整理为知识点概要，并同步生成思维导图，使长视频被转化为可快速浏览的复习结构。",
            5.7,
        ),
        (
            "图 6 重点时间戳与 highlight 片段",
            "该图展示课堂视频分析中的时间定位能力。每个重点片段都包含时间范围、内容概要和标记信息，方便用户在复习或课堂展示时快速跳转到关键位置。",
            5.7,
        ),
        (
            "图 7 产物中心与历史批次管理",
            "该图展示产物中心页面。系统按 run_id 保存每次分析结果，用户可以查看不同任务批次下的摘要、PPT、视频、字幕、导图和报告，避免后续运行覆盖展示版本。",
            5.6,
        ),
        (
            "图 8 运行详情与 Agent 执行轨迹",
            "该图展示运行详情页面。页面记录 Agent 决策、工具调用、校验结果、错误信息和耗时数据，是排查 API 问题、生成质量问题和产物缺失问题的入口。",
            5.6,
        ),
    ]
    new_elements = []
    for image, (caption, note, width) in zip(ui_images, captions):
        add_picture_paragraph(doc, image, width)
        new_elements.append(doc.paragraphs[-1]._p)
        add_caption(doc, caption)
        new_elements.append(doc.paragraphs[-1]._p)
        add_note(doc, note)
        new_elements.append(doc.paragraphs[-1]._p)
    insert_elements_at(doc, insert_idx, new_elements)


def main() -> None:
    docx = final_docx_path()
    doc = Document(docx)
    ui_images = extract_ui_images(doc)
    if len(ui_images) < 7:
        raise RuntimeError(f"Expected at least 7 UI images embedded after diagrams, found {len(ui_images)}")
    insert_idx = remove_old_ui_block(doc)
    build_new_ui_block(doc, ui_images[:7], insert_idx)
    doc.save(docx)
    print(docx)


if __name__ == "__main__":
    main()
